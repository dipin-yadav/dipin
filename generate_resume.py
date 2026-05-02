#!/usr/bin/env python3
"""
Generate an ATS-friendly DOCX resume from portfolio data.
Uses python-docx library for clean, professional formatting.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants


def add_hyperlink(paragraph, url, text, bold=False):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set font properties
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')  # 10pt
    rPr.append(sz)
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    
    run.append(rPr)
    
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)
    
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    
    return hyperlink


def add_bullet_paragraph(doc, text, bold_parts=None):
    """Add a bullet point with optional bold text - compact spacing."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    
    if bold_parts:
        remaining = text
        for bold_text in bold_parts:
            parts = remaining.split(bold_text, 1)
            if len(parts) == 2:
                if parts[0]:
                    run = p.add_run(parts[0])
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9.5)
                run = p.add_run(bold_text)
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                run.font.bold = True
                remaining = parts[1]
        if remaining:
            run = p.add_run(remaining)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
    
    return p


def add_compact_heading(doc, text):
    """Add a compact section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Add underline
    run.font.underline = True
    return p


def add_job_header(doc, title, company, date):
    """Add compact job header."""
    # Title and company on one line
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    run = p.add_run(f' | {company}')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    
    # Date
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    run = p2.add_run(date)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)


def create_resume():
    """Create a compact, one-page ATS-friendly resume with hyperlinks."""
    doc = Document()
    
    # Tight margins for one-page fit
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # ==================== HEADER ====================
    # Name
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(2)
    run = name.add_run('DIPIN YADAV')
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run('Software Developer | Python Backend | Cloud & DevOps')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Contact info with hyperlinks
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(1)
    run = contact.add_run('Gurgaon, Haryana | ')
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    add_hyperlink(contact, 'mailto:dipin.yadav18@gmail.com', 'dipin.yadav18@gmail.com')
    
    # Links row
    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.paragraph_format.space_after = Pt(4)
    run = links.add_run('LinkedIn: ')
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    add_hyperlink(links, 'https://www.linkedin.com/in/dipin-yadav/', 'linkedin.com/in/dipin-yadav')
    run = links.add_run(' | GitHub: ')
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    add_hyperlink(links, 'https://github.com/dipin-yadav', 'github.com/dipin-yadav')
    
    # ==================== PROFESSIONAL SUMMARY (SHORT) ====================
    add_compact_heading(doc, 'PROFESSIONAL SUMMARY')
    
    summary = doc.add_paragraph()
    summary.paragraph_format.space_after = Pt(3)
    run = summary.add_run(
        '4+ years building scalable backend systems with Django, Flask, and Microservices. '
        'AWS Certified Cloud Practitioner. Expert in cloud migrations, Kubernetes, and database optimization. '
        'Track record of architecting high-revenue modules and leading cross-functional teams.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(9.5)
    
    # ==================== CORE COMPETENCIES ====================
    add_compact_heading(doc, 'CORE COMPETENCIES')
    
    skills_para = doc.add_paragraph()
    skills_para.paragraph_format.space_after = Pt(3)
    run = skills_para.add_run(
        'Languages: Python, PHP, SQL, JavaScript, Perl, Bash • '
        'Frameworks: Django, Flask, DRF, Laravel, CakePHP • '
        'Cloud/DevOps: AWS (Certified), Azure, Kubernetes, Docker, CI/CD • '
        'Databases: PostgreSQL, MySQL, MongoDB, Redis'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    
    # ==================== PROFESSIONAL EXPERIENCE ====================
    add_compact_heading(doc, 'PROFESSIONAL EXPERIENCE')
    
    # Job 1
    add_job_header(doc, 'Software Developer (Orange Business)', 'NetConnectGlobal', 'April 2025 – Present')
    job1_bullets = [
        ('Developed "e-files" service replacing legacy FTP/SFTP/SCP protocols across Acort services', ['"e-files"']),
        ('Manage cloud-native apps using Kubernetes on Azure/AWS; migrated SQLite to MySQL', ['Kubernetes', 'SQLite to MySQL']),
        ('Lead Azure cloud migration initiatives; build microservices in Python/Bash/Perl', [])
    ]
    for text, bold_parts in job1_bullets:
        add_bullet_paragraph(doc, text, bold_parts)
    
    # Job 2
    add_job_header(doc, 'Senior Software Engineer', 'Meritto (NoPaperForms)', 'March 2024 – February 2025')
    job2_bullets = [
        ('Built "Prospecta Module" from scratch using Django/Microservices—created new revenue stream', ['"Prospecta Module"', 'Django/Microservices']),
        ('Enhanced Telephony Module with CRM features; designed REST APIs with RBAC', ['REST APIs']),
        ('Integrated third-party APIs (M-cube, Smart-ping) for calling features', [])
    ]
    for text, bold_parts in job2_bullets:
        add_bullet_paragraph(doc, text, bold_parts)
    
    # Job 3
    add_job_header(doc, 'PHP Web Developer', 'Shipway by Unicommerce', 'August 2021 – March 2024')
    job3_bullets = [
        ('Architected Shipway Aggregation module—highest revenue-generating module in company', ['Shipway Aggregation']),
        ('Built Weight MIS, COD Remittance, Wallet systems improving user retention', ['Weight MIS, COD Remittance, Wallet']),
        ('Integrated Shopify, Magento, WooCommerce APIs for e-commerce platforms', [])
    ]
    for text, bold_parts in job3_bullets:
        add_bullet_paragraph(doc, text, bold_parts)
    
    # ==================== CERTIFICATIONS (Compact with links) ====================
    add_compact_heading(doc, 'CERTIFICATIONS')
    
    cert_para = doc.add_paragraph()
    cert_para.paragraph_format.space_after = Pt(3)
    
    certs = [
        ('AWS Cloud Practitioner', 'https://www.linkedin.com/in/dipin-yadav/overlay/Certifications/1543555820/treasury/'),
        ('Quantum Cryptography', 'https://www.linkedin.com/in/dipin-yadav/overlay/Certifications/2104190597/treasury/'),
        ('Fortinet Cybersecurity', 'https://www.credly.com/badges/1473f209-57fb-4073-8861-713bdd1c0485/public_url'),
        ('AI for Beginners', 'https://www.life-global.org/certificate/9b37994d-33ff-4e2d-9124-df1d93bb4313')
    ]
    
    for i, (name, url) in enumerate(certs):
        if i > 0:
            run = cert_para.add_run(' • ')
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
        add_hyperlink(cert_para, url, name)
    
    run = cert_para.add_run(' + 11 more in Cybersecurity & Cloud')
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.italic = True
    
    # ==================== KEY PROJECTS ====================
    add_compact_heading(doc, 'KEY PROJECTS')
    
    projects = [
        ('Orange "e-files" Service', 'Python/Bash/Perl/K8s', 
         'Architected secure internal file sharing service replacing legacy FTP/SFTP/SCP protocols. Deployed on Kubernetes across Azure and AWS with CI/CD pipelines.'),
        ('Azure Cloud Migration', 'Azure/AKS/Docker/Terraform', 
         'Leading strategic migration of cloud-native applications to Azure Kubernetes Service. Optimized infrastructure costs and performance through container orchestration.'),
        ('Acort Database Migration', 'MySQL/SQLite/Python', 
         'Led migration from SQLite to MySQL for enterprise-scale data support. Developed custom ETL scripts ensuring zero data loss and improved query performance.'),
        ('Prospecta Module', 'Django/DRF/PostgreSQL/Microservices', 
         'Built "Prospecta (Raw Data) Module" from scratch using Django microservices. Created new revenue stream with REST APIs and role-based access controls.'),
        ('Shipway Aggregation', 'PHP/MySQL/MongoDB/jQuery', 
         'Developed company\'s highest revenue-generating module. Implemented Weight MIS, COD Remittance, and Wallet systems improving user retention significantly.')
    ]
    
    for i, (name, tech, desc) in enumerate(projects):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f'{name} ')
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.bold = True
        run = p.add_run(f'({tech}): ')
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.italic = True
        run = p.add_run(desc)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
    
    # ==================== EDUCATION ====================
    add_compact_heading(doc, 'EDUCATION')
    
    edu = doc.add_paragraph()
    edu.paragraph_format.space_after = Pt(1)
    run = edu.add_run('B.Tech in Computer Science — Gurgaon Institute of Technology and Management, 2014-2018')
    run.font.name = 'Calibri'
    run.font.size = Pt(9.5)
    run.font.bold = True
    
    # Save the document
    output_path = 'public/Dipin-Yadav-4-years-software-engineer.docx'
    doc.save(output_path)
    print(f"Compact one-page resume generated: {output_path}")
    return output_path


if __name__ == '__main__':
    create_resume()
